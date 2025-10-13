"""
Document converter for .docx files with sentence-level tracking.
Converts DOCX to HTML format with sentence IDs for highlighting.
"""

import re
from pathlib import Path
from zipfile import ZipFile
from typing import List, Dict, Any
import xml.etree.ElementTree as ET

from docx import Document


def read_docx_text(path: Path) -> str:
    """Extract clean visible text from DOCX, including tables, without XML artifacts."""
    doc = Document(path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract text inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    # Join all parts and clean up bullets and hidden chars
    text = "\n".join(parts)
    text = re.sub(r"[•·●▪◦‣–—−]", "-", text)
    text = re.sub(r"[\u0000-\u001F\u2028\u2029\uF020-\uF0FF]", " ", text)
    text = re.sub(r"\s{2,}", " ", text).strip()

    return text


def normalize_sentence(s: str) -> str:
    """Normalize sentence for duplicate comparison."""
    s = s.lower()
    s = s.replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')
    s = re.sub(r"[•·●▪◦‣–—−]", "-", s)  # normalize bullets/dashes
    s = re.sub(r"[^\w\s.,;:()/-]", " ", s)  # strip stray symbols (arrows, boxes, etc.)
    s = re.sub(r"\s{2,}", " ", s).strip()
    return s


def sentence_split(text: str) -> List[str]:
    """Naive sentence split with extra break for overly long lines."""
    t = re.sub(r"[\r\n]+", " ", text)
    # Replace list bullets/newlines before splitting
    t = re.sub(r"[•·●▪◦‣]", ".", t)  # treat bullets as sentence boundaries
    t = re.sub(r"[\r\n]+", " ", t)
    parts = re.split(r"(?<=[\.\!\?\:;])\s+", t)
    out = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        toks = p.split()
        if len(toks) > 80:
            for i in range(0, len(toks), 30):
                out.append(" ".join(toks[i:i+30]))
        else:
            out.append(p)
    return out


def tokenize_words(s: str) -> List[str]:
    """Tokenize sentence into words for similarity checks."""
    s = re.sub(r"[•·●▪◦‣–—−]", " ", s)
    return re.findall(r"[a-z0-9]+", s.lower())


def convert_docx_to_html(docx_path: Path, min_words: int = 8) -> Dict[str, Any]:
    """
    Convert DOCX to HTML with sentence-level IDs.

    Returns:
        {
            "filename": str,
            "html": str,
            "sentences": List[{"id": int, "text": str, "normalized": str}]
        }
    """
    text = read_docx_text(docx_path)
    sentences_raw = sentence_split(text)

    sentences = []
    html_parts = []
    sent_id = 0

    html_parts.append("<div class='document-content'>")

    for raw_sent in sentences_raw:
        norm = normalize_sentence(raw_sent)
        word_count = len(tokenize_words(norm))

        # --- Escape HTML entities safely ---
        escaped_sent = (
            raw_sent.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
        )

        if word_count >= min_words:
            sentences.append({
                "id": sent_id,
                "text": raw_sent,
                "normalized": norm,
                "word_count": word_count
            })

            html_parts.append(
                f'<span class="sentence" data-sentence-id="{sent_id}">{escaped_sent}</span> '
            )
            sent_id += 1
        else:
            html_parts.append(f'<span class="sentence-short">{escaped_sent}</span> ')

    html_parts.append("</div>")

    return {
        "filename": docx_path.name,
        "html": "".join(html_parts),
        "sentences": sentences,
        "total_sentences": len(sentences)
    }


def get_document_structure(docx_path: Path) -> Dict[str, Any]:
    """
    Extract document structure with headings and paragraphs.
    More advanced conversion that preserves document structure.
    """
    try:
        from docx import Document
        
        doc = Document(docx_path)
        structure = {
            "filename": docx_path.name,
            "paragraphs": [],
            "headings": []
        }
        
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                structure["headings"].append({
                    "level": para.style.name,
                    "text": para.text
                })
            structure["paragraphs"].append({
                "text": para.text,
                "style": para.style.name
            })
        
        return structure
    except ImportError:
        # Fallback if python-docx not available
        return {
            "filename": docx_path.name,
            "paragraphs": [{"text": read_docx_text(docx_path), "style": "Normal"}],
            "headings": []
        }

